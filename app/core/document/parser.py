"""Document parsers: PDF (Docling + PyMuPDF fallback), Word (python-docx), TXT, Markdown."""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

import re

from app.config import settings

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════
# Cross-page break repair
# ═══════════════════════════════════════════════════════════════════

def fix_cross_page_breaks(text: str) -> str:
    """Repair text artifacts caused by PDF page boundaries.

    Fixes:
    1. Hyphenation at page break: ``super-\\nvised`` → ``supervised``
    2. Mid-sentence newlines (Chinese & English)
    3. Orphan page numbers
    4. Repeated headers/footers (appear ≥3 times)
    5. Excessive blank lines
    """
    # 1. Hyphenation break: "super-\nvised" → "supervised"
    text = re.sub(r'(\w+)-\n(\w+)', r'\1\2', text)

    # 2. Mid-sentence break — line ends without punctuation,
    #    next line starts with lowercase or Chinese character
    text = re.sub(r'([^。！？.!?\n])\n([a-z一-鿿])', r'\1\2', text)

    # 3. Orphan page-number lines (standalone 1-3 digit numbers)
    text = re.sub(r'\n\d{1,3}\n', '\n', text)

    # 4. Repeated headers/footers — same short line appearing ≥3 times
    lines = text.split('\n')
    freq: dict[str, int] = {}
    for line in lines:
        stripped = line.strip()
        if 2 < len(stripped) < 80:
            freq[stripped] = freq.get(stripped, 0) + 1
    for line, count in freq.items():
        if count >= 3:
            text = text.replace(line + '\n', '')
            text = text.replace('\n' + line, '')
            text = text.replace(line, '')

    # 5. Collapse excessive blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


@dataclass
class ParsedDocument:
    """Parsed document with metadata."""
    text: str
    filename: str
    page_count: int = 1
    metadata: dict | None = None
    # Image-rich pages: (page_number, image_path, page_text)
    image_pages: list[dict] = field(default_factory=list)


# ═══════════════════════════════════════════════════════════════════
# PDF — Docling (layout-aware structured extraction)
# ═══════════════════════════════════════════════════════════════════

def parse_pdf_docling(filepath: Path, extract_images: bool = True) -> ParsedDocument:
    """Extract text from PDF using Docling for layout-aware structured output.

    Docling automatically:
    - Detects heading hierarchy (H1-H6)
    - Extracts tables with TableFormer (preserves merged cells, multi-level headers)
    - Converts formulas to LaTeX
    - Preserves reading order and list structure
    - Handles multi-column layouts

    PyMuPDF is retained for image page detection (VLM pipeline).
    Falls back to parse_pdf_pymupdf if Docling is unavailable or fails.
    """
    import fitz  # pymupdf
    from docling.document_converter import DocumentConverter

    # 1. Docling — structured markdown extraction
    converter = DocumentConverter()
    result = converter.convert(str(filepath))
    structured_text = result.document.export_to_markdown()
    page_count = len(result.pages) if hasattr(result, "pages") else 1

    # Collect table info for metadata
    table_count = 0
    if hasattr(result.document, "tables"):
        table_count = len(result.document.tables)

    logger.info(
        "Docling parsed '%s': %d pages, %d chars, %d tables",
        filepath.name, page_count, len(structured_text), table_count,
    )

    # 2. PyMuPDF — image detection (VLM pipeline)
    image_pages: list[dict] = []
    if extract_images:
        doc = fitz.open(str(filepath))
        for page_num in range(doc.page_count):
            page = doc[page_num]
            image_list = page.get_images(full=True)
            if image_list:
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat)
                img_path = filepath.parent / f"_img_{filepath.stem}_p{page_num + 1}.png"
                pix.save(str(img_path))
                image_pages.append({
                    "page_number": page_num + 1,
                    "image_path": str(img_path),
                    "image_count": len(image_list),
                    "page_text": page.get_text("text").strip()[:500],
                })
        doc.close()

    return ParsedDocument(
        text=structured_text,
        filename=filepath.name,
        page_count=page_count,
        metadata={
            "source": str(filepath),
            "type": "pdf",
            "parser": "docling",
            "tables_detected": table_count,
        },
        image_pages=image_pages,
    )


# ═══════════════════════════════════════════════════════════════════
# PDF — PyMuPDF (plain text fallback)
# ═══════════════════════════════════════════════════════════════════

def parse_pdf_pymupdf(filepath: Path, extract_images: bool = True) -> ParsedDocument:
    """Extract plain text from PDF using PyMuPDF (fitz).

    When extract_images=True, pages with embedded images are rendered
    as PNG files for later VLM analysis.
    """
    import fitz  # pymupdf

    doc = fitz.open(str(filepath))
    num_pages = doc.page_count
    pages: list[str] = []
    image_pages: list[dict] = []

    for page_num in range(num_pages):
        page = doc[page_num]
        text = page.get_text("text")
        if text.strip():
            pages.append(text)

        if extract_images:
            image_list = page.get_images(full=True)
            if image_list:
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat)
                img_path = filepath.parent / f"_img_{filepath.stem}_p{page_num + 1}.png"
                pix.save(str(img_path))
                image_pages.append({
                    "page_number": page_num + 1,
                    "image_path": str(img_path),
                    "image_count": len(image_list),
                    "page_text": text.strip()[:500],
                })

    doc.close()
    full_text = "\n\n".join(pages)
    logger.info(
        "PyMuPDF parsed '%s': %d pages, %d chars, %d image pages",
        filepath.name, num_pages, len(full_text), len(image_pages),
    )
    return ParsedDocument(
        text=full_text,
        filename=filepath.name,
        page_count=num_pages,
        metadata={"source": str(filepath), "type": "pdf", "parser": "pymupdf"},
        image_pages=image_pages,
    )


# ═══════════════════════════════════════════════════════════════════
# Other formats
# ═══════════════════════════════════════════════════════════════════

def parse_docx(filepath: Path) -> ParsedDocument:
    """Extract text from a Word document."""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(filepath))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    logger.info("Parsed DOCX '%s': %d paragraphs, %d chars", filepath.name, len(paragraphs), len(full_text))
    return ParsedDocument(
        text=full_text,
        filename=filepath.name,
        metadata={"source": str(filepath), "type": "docx"},
    )


def parse_txt(filepath: Path) -> ParsedDocument:
    """Read a plain text file."""
    text = filepath.read_text(encoding="utf-8")
    logger.info("Read TXT '%s': %d chars", filepath.name, len(text))
    return ParsedDocument(
        text=text,
        filename=filepath.name,
        metadata={"source": str(filepath), "type": "txt"},
    )


def parse_markdown(filepath: Path) -> ParsedDocument:
    """Read a markdown file."""
    text = filepath.read_text(encoding="utf-8")
    logger.info("Read MD '%s': %d chars", filepath.name, len(text))
    return ParsedDocument(
        text=text,
        filename=filepath.name,
        metadata={"source": str(filepath), "type": "markdown"},
    )


# ═══════════════════════════════════════════════════════════════════
# Dispatcher
# ═══════════════════════════════════════════════════════════════════

PARSERS = {
    ".pdf": "pdf",       # handled by parse_pdf with docling/pymupdf dispatch
    ".docx": parse_docx,
    ".doc": parse_docx,
    ".txt": parse_txt,
    ".md": parse_markdown,
    ".markdown": parse_markdown,
}


def parse_pdf(filepath: Path, extract_images: bool = True) -> ParsedDocument:
    """Parse PDF with configured parser (docling or pymupdf).

    Uses settings.pdf_parser to decide. Falls back to PyMuPDF
    if Docling is unavailable or fails.
    """
    if settings.pdf_parser == "docling":
        try:
            result = parse_pdf_docling(filepath, extract_images)
        except Exception:
            logger.warning(
                "Docling failed for '%s', falling back to PyMuPDF", filepath.name,
                exc_info=True,
            )
            result = parse_pdf_pymupdf(filepath, extract_images)
    else:
        result = parse_pdf_pymupdf(filepath, extract_images)

    # Apply cross-page break repair to clean boundary artifacts
    result.text = fix_cross_page_breaks(result.text)
    return result


def parse_document(filepath: Path, extract_images: bool = True) -> ParsedDocument:
    """Parse any supported document type."""
    ext = filepath.suffix.lower()
    parser = PARSERS.get(ext)
    if parser is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {list(PARSERS.keys())}")
    if ext == ".pdf":
        return parse_pdf(filepath, extract_images=extract_images)
    return parser(filepath)
